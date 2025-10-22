# literature/utils.py
import io
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.styles import ParagraphStyle
from docx import Document


def export_review_to_pdf(review_text: str, topic: str) -> bytes:
    """
    Convert a long review text into a styled PDF document and return bytes.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    # Define a justified paragraph style
    justified_style = ParagraphStyle(
        "Justify",
        parent=styles["Normal"],
        alignment=TA_JUSTIFY,
        fontSize=11,
        leading=14
    )

    story = []
    story.append(Paragraph(f"<b>Literature Review: {topic}</b>", styles["Title"]))
    story.append(Spacer(1, 0.3 * inch))

    # Split text by double newlines to keep section breaks
    for section in review_text.split("\n\n"):
        if section.strip():
            story.append(Paragraph(section.strip(), justified_style))
            story.append(Spacer(1, 0.15 * inch))

    doc.build(story)
    pdf_data = buffer.getvalue()
    buffer.close()
    return pdf_data


def export_review_to_docx(review_text: str, topic: str) -> bytes:
    """
    Convert the review text into a .docx Word file and return bytes.
    """
    doc = Document()
    doc.add_heading(f"Literature Review: {topic}", level=1)

    for section in review_text.split("\n\n"):
        if section.strip():
            doc.add_paragraph(section.strip(), style="Normal")

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
